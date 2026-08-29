"""
Tests for learning a .wvfmt from a .docx (wordvault/printing/
learn_format.py): build a Word file with known geometry and styles,
learn from it, and the resulting format must load cleanly and print
the same numbers back.
"""

import zipfile

import pytest

docx = pytest.importorskip("docx")

from wordvault.printing.format_file import load_format  # noqa: E402
from wordvault.printing.learn_format import learn_format  # noqa: E402


def make_sample(path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    d = docx.Document()
    section = d.sections[0]
    section.page_width = Inches(6)          # the KDP trim
    section.page_height = Inches(9)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(0.7)

    normal = d.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(10)

    h1 = d.styles["Heading 1"]
    h1.font.size = Pt(17)
    h1.font.bold = True

    d.add_heading("Chapter", level=1)
    d.add_paragraph("Body words.")
    d.save(str(path))


def test_learned_format_loads_and_matches(tmp_path):
    sample = tmp_path / "sample.docx"
    make_sample(sample)

    toml_text = learn_format(sample, "Learned Look")
    target = tmp_path / "learned.wvfmt"
    target.write_text(toml_text, encoding="utf-8")
    fmt = load_format(target)               # must validate cleanly

    assert fmt.name == "Learned Look"
    assert fmt.page_size == "6x9"           # recognized by dimensions
    top, right, bottom, left = fmt.margins.for_page(0)
    assert top == pytest.approx(0.7 * 25.4, abs=0.5)
    assert bottom == pytest.approx(0.5 * 25.4, abs=0.5)
    assert left == pytest.approx(1.2 * 25.4, abs=0.5)
    assert fmt.body.font == "Georgia"
    assert fmt.body.size_pt == pytest.approx(11)
    assert fmt.body.align == "justify"
    assert fmt.body.space_after_pt == pytest.approx(10, abs=0.5)
    h1 = fmt.style_for_heading(1)
    assert h1.size_pt == pytest.approx(17)
    assert h1.bold
    # The file says where it came from.
    assert "learned from sample.docx" in toml_text


def test_direct_formatting_outvotes_the_style_sheet(tmp_path):
    """Andrew's Inside_Gods_House8.docx: the style sheet said Times
    New Roman 12 (docDefaults from an ancestral template, Normal style
    silent) while every visible word wore direct-formatted Aptos 11
    with 12pt space-after painted on the paragraphs.  Word shows the
    paint, so the learner must learn the paint."""
    from docx.shared import Pt

    d = docx.Document()
    # The undercoat: document defaults claim Times New Roman 12.
    styles_el = d.styles.element
    rpr_default = styles_el.xpath("w:docDefaults/w:rPrDefault/w:rPr")[0]
    ns = ("{http://schemas.openxmlformats.org/wordprocessingml/2006/"
          "main}")
    for el in list(rpr_default):
        rpr_default.remove(el)
    fonts = rpr_default.makeelement(f"{ns}rFonts", {f"{ns}ascii":
                                                    "Times New Roman"})
    size = rpr_default.makeelement(f"{ns}sz", {f"{ns}val": "24"})
    rpr_default.append(fonts)
    rpr_default.append(size)

    # The paint: every body run direct-formatted Aptos 11, paragraphs
    # given 12pt space-after directly.
    for i in range(6):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"Directly formatted words number {i}.")
        run.font.name = "Aptos"
        run.font.size = Pt(11)
    d.save(str(tmp_path / "painted.docx"))

    toml_text = learn_format(tmp_path / "painted.docx", "Painted")
    target = tmp_path / "painted.wvfmt"
    target.write_text(toml_text, encoding="utf-8")
    fmt = load_format(target)
    assert fmt.body.font == "Aptos"
    assert fmt.body.size_pt == pytest.approx(11)
    assert fmt.body.space_after_pt == pytest.approx(12, abs=0.5)


def test_paragraph_defaults_supply_the_gap(tmp_path):
    """Modern Word keeps its standard paragraph spacing in
    docDefaults/pPrDefault, not in the Normal style.  A learner that
    misses it emits a format with NO paragraph gap — Andrew's 'no
    full line between paragraphs' printed page."""
    d = docx.Document()
    # Strip the template Normal style's own spacing so the defaults
    # are the only source (when Normal DOES say spacing, it correctly
    # outranks pPrDefault — that's Word's precedence).
    normal = d.styles["Normal"].paragraph_format
    normal.space_after = None
    normal.line_spacing = None
    ns = ("{http://schemas.openxmlformats.org/wordprocessingml/2006/"
          "main}")
    dd = d.styles.element.find(f"{ns}docDefaults")
    spacing = dd.find(f"{ns}pPrDefault/{ns}pPr/{ns}spacing")
    if spacing is None:                  # template variations
        ppr_default = dd.find(f"{ns}pPrDefault")
        if ppr_default is None:
            ppr_default = dd.makeelement(f"{ns}pPrDefault", {})
            dd.append(ppr_default)
        ppr = ppr_default.find(f"{ns}pPr")
        if ppr is None:
            ppr = ppr_default.makeelement(f"{ns}pPr", {})
            ppr_default.append(ppr)
        spacing = ppr.makeelement(f"{ns}spacing", {})
        ppr.append(spacing)
    spacing.set(f"{ns}after", "160")
    spacing.set(f"{ns}line", "278")
    spacing.set(f"{ns}lineRule", "auto")
    d.add_paragraph("Body words that carry no direct formatting.")
    sample = tmp_path / "defaults.docx"
    d.save(str(sample))

    target = tmp_path / "defaults.wvfmt"
    target.write_text(learn_format(sample, "Defaults"),
                      encoding="utf-8")
    fmt = load_format(target)
    assert fmt.body.space_after_pt == pytest.approx(8, abs=0.1)
    assert fmt.body.line_spacing == pytest.approx(278 / 240, abs=0.01)


def test_page_number_footer_is_detected(tmp_path):
    """A footer with Word's PAGE field becomes [footer] {page}."""
    sample = tmp_path / "footed.docx"
    make_sample(sample)
    # Splice a PAGE-field footer into the zip by hand (python-docx
    # cannot author fields).
    footer_xml = (
        '<?xml version="1.0"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main"><w:p><w:r>'
        '<w:instrText> PAGE </w:instrText></w:r></w:p></w:ftr>')
    with zipfile.ZipFile(sample, "a") as zf:
        zf.writestr("word/footer1.xml", footer_xml)

    toml_text = learn_format(sample, "Footed")
    assert '[footer]' in toml_text and '"{page}"' in toml_text
    target = tmp_path / "footed.wvfmt"
    target.write_text(toml_text, encoding="utf-8")
    assert load_format(target).footer.center == "{page}"
