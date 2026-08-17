"""
End-to-end tests for the ingest pipeline (Phases A and B).

Builds real .docx files in a temp folder with python-docx, then runs the
same Ingestor the CLI uses.  Skipped automatically when python-docx is
not installed (the rest of WordVault never requires it).
"""

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from wordvault import DocumentStore  # noqa: E402
from wordvault.ingest import Ingestor  # noqa: E402
from wordvault.ingest.extract import extract_text, normalize_text  # noqa: E402

ESSAY = (
    "In the beginning God created the heaven and the earth. And the earth "
    "was without form, and void; and darkness was upon the face of the deep. "
) * 20

OTHER = (
    "The quarterly report shows revenue growth across all divisions with "
    "particular strength in the northern region this year. "
) * 20


def make_docx(path, paragraphs):
    """Write a real .docx file with the given paragraph texts."""
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(path))


@pytest.fixture()
def library(tmp_path):
    """A source folder of five .docx files and an empty library store:
    draft1/draft2 are versions, copy is an exact duplicate of draft1,
    unrelated stands alone, empty has no text."""
    src = tmp_path / "src"
    src.mkdir()
    make_docx(src / "essay_draft1.docx", [ESSAY])
    make_docx(src / "essay_draft2.docx",
              [ESSAY.replace("darkness", "the dark") + " A new closing thought."])
    make_docx(src / "essay_copy_other_name.docx", [ESSAY])  # exact dup of draft1
    make_docx(src / "unrelated.docx", [OTHER])
    make_docx(src / "empty.docx", [])

    store = DocumentStore(tmp_path / "library.db")
    yield src, store
    store.close()


def test_extract_normalizes(tmp_path):
    p = tmp_path / "one.docx"
    make_docx(p, ["Line one  ", "", "Line two"])   # trailing spaces stripped
    assert extract_text(p) == "Line one\n\nLine two\n"
    assert normalize_text("a\r\nb\r") == "a\nb\n"


def test_style_spacing_becomes_blank_lines(tmp_path):
    # Word shows space between paragraphs via "space after" settings,
    # with no empty paragraphs anywhere. The extractor must reproduce
    # that visual space as blank lines.
    from docx.shared import Pt

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    a = d.add_paragraph("First paragraph.")
    a.paragraph_format.space_after = Pt(8)      # visually spaced in Word
    d.add_paragraph("Second paragraph.")
    path = tmp_path / "spaced.docx"
    d.save(str(path))

    assert extract_markdown(path) == "First paragraph.\n\nSecond paragraph.\n"


def test_paragraphs_always_get_a_blank_line(tmp_path):
    # The separation rule (Aug 2026, by request): EVERY pair of
    # neighboring paragraphs is separated by a blank line — even when
    # Word's spacing was an explicit 0pt.  Predictable beats faithful.
    from docx.shared import Pt

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    for text in ["ISBN: 123", "Published in the USA"]:
        p = d.add_paragraph(text)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    path = tmp_path / "tight.docx"
    d.save(str(path))

    assert extract_markdown(path) == "ISBN: 123\n\nPublished in the USA\n"


def test_headings_always_get_breathing_room(tmp_path):
    from docx.shared import Pt

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    p = d.add_paragraph("Intro text.")
    p.paragraph_format.space_after = Pt(0)      # even with tight spacing...
    d.add_paragraph("Chapter One", style="Heading 1")
    q = d.add_paragraph("Chapter text.")
    q.paragraph_format.space_before = Pt(0)
    path = tmp_path / "headed.docx"
    d.save(str(path))

    # ...headings are separated on both sides anyway.
    assert extract_markdown(path) == (
        "Intro text.\n\n# Chapter One\n\nChapter text.\n"
    )


def test_normalize_caps_blank_line_runs():
    # A Word title page's dozen empty paragraphs must not become a wall
    # of blank space — runs collapse to at most two blank lines.
    walled = "Title\n" + "\n" * 14 + "Author\n\nBody text\n"
    assert normalize_text(walled) == "Title\n\n\nAuthor\n\nBody text\n"
    # Single blank lines (normal paragraphing) pass through untouched.
    assert normalize_text("one\n\ntwo\n") == "one\n\ntwo\n"


def test_extract_markdown_carries_formatting(tmp_path):
    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    d.add_paragraph("The Coming Kingdom", style="Heading 1")
    d.add_paragraph("First Signs", style="Heading 2")
    p = d.add_paragraph()
    p.add_run("The word ")
    p.add_run("kingdom").bold = True
    p.add_run(" appears ")
    run = p.add_run("often")
    run.italic = True
    p.add_run(" here.")
    d.add_paragraph("a marked point", style="List Bullet")
    path = tmp_path / "formatted.docx"
    d.save(str(path))

    text = extract_markdown(path)
    assert "# The Coming Kingdom\n" in text
    assert "## First Signs\n" in text
    assert "The word **kingdom** appears *often* here." in text
    assert "- a marked point\n" in text


def test_extract_markdown_merges_split_runs(tmp_path):
    # Word often splits one bold phrase into several runs; the extractor
    # must merge them into ONE pair of markers.
    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    p = d.add_paragraph()
    for piece in ("all ", "of ", "this bold"):
        r = p.add_run(piece)
        r.bold = True
    path = tmp_path / "runs.docx"
    d.save(str(path))

    assert extract_markdown(path).strip() == "**all of this bold**"


def test_toolbar_list_imports_as_bullets(tmp_path):
    """Lists made with the RIBBON BUTTONS (numbering attached directly
    to ordinary paragraphs, style still 'Normal') must import as list
    items — they used to flatten into plain prose."""
    from docx.oxml.ns import qn

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    d.add_paragraph("Shopping:")
    for item in ("bread", "wine"):
        p = d.add_paragraph(item)
        numpr = p._p.get_or_add_pPr().get_or_add_numPr()
        numpr.get_or_add_numId().val = 7      # a numbering reference...
        numpr.get_or_add_ilvl().val = 0
        assert p._p.pPr.numPr.numId.val == 7
        _ = qn  # (namespace helper imported for clarity of intent)
    path = tmp_path / "toolbar.docx"
    d.save(str(path))

    import re

    text = extract_markdown(path)
    # Bullet or number depends on what numbering.xml says numId 7 is
    # (the classifier reads the real definitions — python-docx's own
    # template defines it as decimal).  Either way: LIST items, tight.
    assert re.search(r"Shopping:\n\n(- |1\. )bread\n\1wine\n", text), text


def test_numbering_formats_classifier(tmp_path):
    """word/numbering.xml decides bullet vs number per numId."""
    import zipfile

    from wordvault.ingest.extract import _numbering_formats

    xml = (
        '<w:numbering xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">'
        '<w:abstractNum w:abstractNumId="0">'
        '<w:lvl w:ilvl="0"><w:numFmt w:val="bullet"/></w:lvl></w:abstractNum>'
        '<w:abstractNum w:abstractNumId="1">'
        '<w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl></w:abstractNum>'
        '<w:num w:numId="5"><w:abstractNumId w:val="0"/></w:num>'
        '<w:num w:numId="6"><w:abstractNumId w:val="1"/></w:num>'
        '</w:numbering>'
    )
    path = tmp_path / "n.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("word/numbering.xml", xml)
    formats = _numbering_formats(path)
    assert formats[(5, 0)] == "bullet"
    assert formats[(6, 0)] == "number"


def test_hyperlink_text_survives_with_address(tmp_path):
    """Linked words used to VANISH (python-docx's runs omit hyperlink
    contents).  Now: 'the text (address)'."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    p = d.add_paragraph("See ")
    rid = d.part.relate_to("https://example.com/study",
                           RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "the article"
    run.append(text_el)
    link.append(run)
    p._p.append(link)
    path = tmp_path / "linked.docx"
    d.save(str(path))

    assert ("See the article (https://example.com/study)"
            in extract_markdown(path))


def test_table_text_flattens_to_rows(tmp_path):
    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    d.add_paragraph("Before the table.")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feast"
    table.cell(0, 1).text = "Month"
    table.cell(1, 0).text = "Passover"
    table.cell(1, 1).text = "First"
    d.add_paragraph("After the table.")
    path = tmp_path / "tabled.docx"
    d.save(str(path))

    text = extract_markdown(path)
    assert "Before the table.\n\nFeast — Month\nPassover — First\n\n" \
        "After the table." in text


def test_underline_becomes_italic_and_indent_becomes_quote(tmp_path):
    from docx.shared import Inches

    from wordvault.ingest.extract import extract_markdown

    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("a ")
    p.add_run("stressed").underline = True
    p.add_run(" word")
    q = d.add_paragraph("Blessed are the meek, for they shall "
                        "inherit the earth.")
    q.paragraph_format.left_indent = Inches(0.5)
    path = tmp_path / "ui.docx"
    d.save(str(path))

    text = extract_markdown(path)
    assert "a *stressed* word" in text
    assert "> Blessed are the meek" in text


def test_word_internal_dates_beat_filesystem(tmp_path):
    """The dates INSIDE the docx (Word's own record of authoring) win
    over filesystem dates, which every copy or sync resets."""
    from datetime import datetime, timezone

    from wordvault.ingest.extract import (
        docx_internal_dates,
        document_dates_utc,
    )

    d = docx.Document()
    d.add_paragraph("Old words.")
    d.core_properties.created = datetime(2019, 3, 4, 16, 30,
                                         tzinfo=timezone.utc)
    d.core_properties.modified = datetime(2021, 7, 1, 9, 0,
                                          tzinfo=timezone.utc)
    path = tmp_path / "aged.docx"
    d.save(str(path))       # the FILE is created right now (copy-day)

    created, modified = docx_internal_dates(path)
    assert created.startswith("2019-03-04")
    assert modified.startswith("2021-07-01")
    # The combined best-evidence reader prefers the internal record.
    best_created, best_modified = document_dates_utc(path)
    assert best_created.startswith("2019-03-04")
    assert best_modified.startswith("2021-07-01")


def test_repair_document_dates_fixes_copy_day_metadata(tmp_path):
    """The Refresh verification: a document imported with copy-day
    dates gets them corrected to the Word file's internal record —
    once; a second pass finds nothing to fix."""
    from datetime import datetime, timezone

    from wordvault.ingest.pipeline import repair_document_dates

    d = docx.Document()
    d.add_paragraph("Words with a past.")
    d.core_properties.created = datetime(2018, 1, 2, tzinfo=timezone.utc)
    d.core_properties.modified = datetime(2020, 5, 6, tzinfo=timezone.utc)
    path = tmp_path / "past.docx"
    d.save(str(path))

    store = DocumentStore(tmp_path / "lib.db")
    doc = store.create_document(
        "Past", original_path=str(path),
        original_mtime="2026-08-16T00:00:00+00:00",   # copy-day lies
        created_utc="2026-08-16T00:00:00+00:00",
    )
    assert repair_document_dates(store, doc) is True
    fixed = store.get_document(doc.id)
    assert fixed.created_utc.startswith("2018-01-02")
    assert fixed.original_mtime.startswith("2020-05-06")
    # Second pass: everything already agrees.
    assert repair_document_dates(store, fixed) is False
    store.close()


def test_phase_a_ingests_and_collapses_duplicates(library):
    src, store = library
    stats = Ingestor(store).ingest_folder(src)

    assert stats.files_seen == 5
    assert stats.ingested == 3        # draft1, draft2, unrelated
    assert stats.duplicates == 1      # the exact copy under another name
    assert stats.empty == 1
    assert not stats.errors

    docs = store.ingested_documents()
    assert len(docs) == 3
    # Documents carry their file's path and an ingest-origin revision.
    for doc in docs:
        revs = store.list_revisions(doc.id)
        assert len(revs) == 1 and revs[0].origin == "ingest"

    # The duplicate's path was noted on the surviving document.
    dup_owner = next(d for d in docs if "draft1" in d.original_path)
    assert any("copy_other_name" in p for p in store.ingest_duplicates_for(dup_owner.id))


def test_phase_b_proposes_version_group(library):
    src, store = library
    stats = Ingestor(store, threshold=0.5).ingest_folder(src)

    assert stats.groups_proposed == 1
    gid = store.list_similarity_groups("pending")[0]
    titles = sorted(d.title for d, _ in store.group_members(gid))
    assert titles == ["essay_draft1", "essay_draft2"]   # unrelated stays out


def test_rerun_is_idempotent(library):
    src, store = library
    Ingestor(store).ingest_folder(src)
    stats2 = Ingestor(store).ingest_folder(src)   # run again, same folder

    assert stats2.ingested == 0
    assert stats2.duplicates == 0                 # dup not re-noted either
    assert stats2.skipped_known == 4              # 3 documents + 1 dup path
    assert len(store.ingested_documents()) == 3


def test_limit_allows_trial_run(library):
    src, store = library
    stats = Ingestor(store).ingest_folder(src, limit=1)
    assert stats.ingested == 1
    # Continuing without the limit picks up the rest.
    stats2 = Ingestor(store).ingest_folder(src)
    assert len(store.ingested_documents()) == 3


def test_archive_copies_ingested_files(library, tmp_path):
    src, store = library
    archive = tmp_path / "originals"
    ticks = []
    stats = Ingestor(
        store, archive_dir=archive, tick=lambda: ticks.append(1)
    ).ingest_folder(src)

    # One copy per NEW document (not for duplicates or empty files),
    # named "<doc-id> - <filename>" so any document's source is findable.
    assert stats.archived == stats.ingested == 3
    copies = sorted(p.name for p in archive.iterdir())
    assert len(copies) == 3
    for doc in store.ingested_documents():
        assert any(name == f"{doc.id:05d} - {Path(doc.original_path).name}"
                   for name in copies)
    # The tick callback fired for every file seen (GUI progress heartbeat).
    assert len(ticks) == stats.files_seen


def test_archive_skips_already_known_on_rerun(library, tmp_path):
    src, store = library
    archive = tmp_path / "originals"
    Ingestor(store, archive_dir=archive).ingest_folder(src)
    stats2 = Ingestor(store, archive_dir=archive).ingest_folder(src)
    assert stats2.archived == 0                    # nothing new, no re-copy
    assert len(list(archive.iterdir())) == 3


def test_confirmed_groups_survive_rerun(library):
    src, store = library
    Ingestor(store, threshold=0.5).ingest_folder(src)
    gid = store.list_similarity_groups("pending")[0]
    store.set_group_status(gid, "confirmed")      # the author decided

    Ingestor(store, threshold=0.5).ingest_folder(src)   # detector re-runs
    assert store.list_similarity_groups("confirmed") == [gid]
    assert store.list_similarity_groups("pending") == []  # not re-proposed
