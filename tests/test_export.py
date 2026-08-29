"""
Tests for Markdown -> .docx export (wordvault/export_docx.py).

The crown test is the ROUND TRIP: export through markdown_to_docx,
re-import through the ingest extractor, and the structure must
survive the full circle — headings, emphasis, lists, quotes,
paragraph joins.  Exporter and importer are exact inverses, and this
is where that claim is enforced.
"""

import pytest

docx = pytest.importorskip("docx")

from wordvault.export_docx import markdown_to_docx  # noqa: E402
from wordvault.ingest.extract import extract_markdown  # noqa: E402

MARKDOWN = """# The Coming Kingdom

## First Signs

The word **kingdom** appears *often* and ***emphatically*** here.

A second paragraph, whose two source lines
join into one, as the conventions say.

> A quoted verse stands apart.

- first point
- second point

1. alpha
2. beta
"""


def test_round_trip_survives(tmp_path):
    path = tmp_path / "out.docx"
    markdown_to_docx(MARKDOWN, path, title="The Coming Kingdom",
                     author="Andrew Hopkins")
    back = extract_markdown(path)

    assert "# The Coming Kingdom" in back
    assert "## First Signs" in back
    assert "**kingdom**" in back
    assert "*often*" in back
    assert "***emphatically***" in back
    assert "join into one, as the conventions say." in back
    assert "whose two source lines join" in back      # joined, not split
    assert "> A quoted verse stands apart." in back
    assert "- first point" in back and "- second point" in back
    assert "1. alpha" in back                          # a numbered list
    # And the Word file's own identity fields carry the metadata.
    d = docx.Document(str(path))
    assert d.core_properties.title == "The Coming Kingdom"
    assert d.core_properties.author == "Andrew Hopkins"


def test_unmatched_markers_stay_literal(tmp_path):
    path = tmp_path / "lit.docx"
    markdown_to_docx("A lone *asterisk pair left open.\n", path)
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "*asterisk" in text          # untouched, exactly as typed


def test_heading_levels_map_to_word_styles(tmp_path):
    path = tmp_path / "h.docx"
    markdown_to_docx("# One\n\n### Three\n\nBody.\n", path)
    d = docx.Document(str(path))
    styles = [p.style.name for p in d.paragraphs if p.text]
    assert "Heading 1" in styles and "Heading 3" in styles


def test_pipe_tables_become_real_word_tables(tmp_path):
    """The provenance report's sessions table must arrive in Word as a
    real table (header row bold, separator row dropped) — not a wall
    of pipe characters."""
    import docx as docx_lib

    from wordvault.export_docx import markdown_to_docx

    md = ("# Report\n\n"
          "| Session | Began | Net |\n"
          "|---|---|---|\n"
          "| 1 | 2026-08-01 | +400 |\n"
          "| 2 | 2026-08-03 | -20 |\n\n"
          "Prose after the table.\n")
    path = tmp_path / "table.docx"
    markdown_to_docx(md, path, title="T", author="A")

    d = docx_lib.Document(str(path))
    assert len(d.tables) == 1
    table = d.tables[0]
    assert len(table.rows) == 3            # header + 2 data rows
    assert table.cell(0, 0).text == "Session"
    assert table.cell(2, 2).text == "-20"
    assert "|" not in "\n".join(p.text for p in d.paragraphs)
    assert any("Prose after the table." in p.text for p in d.paragraphs)


def test_compact_mode_wears_andrews_dress(tmp_path):
    """The provenance report's Word dress, learned from Andrew's
    hand-tuned file: 0.4-inch margins, 8 pt everywhere (table cells
    included), single spacing — and the corrections section starts a
    fresh page."""
    import docx as docx_lib
    from docx.shared import Inches, Pt

    from wordvault.export_docx import markdown_to_docx

    md = ("# Provenance Report — Essay\n\n"
          "The document facts.\n\n"
          "| Session | Net |\n|---|---|\n| 1 | +400 |\n\n"
          "## Corrections along the way\n\n"
          "teh → the;  recieve → receive\n")
    path = tmp_path / "compact.docx"
    markdown_to_docx(md, path, title="T", author="A", compact=True,
                     page_break_before=("Corrections along the way",))

    d = docx_lib.Document(str(path))
    section = d.sections[0]
    assert section.top_margin == Inches(0.4)
    assert section.left_margin == Inches(0.4)
    sized = [r.font.size for p in d.paragraphs for r in p.runs
             if r.text.strip()]
    assert sized and all(s == Pt(8) for s in sized)
    cell_runs = [r.font.size for t in d.tables for row in t.rows
                 for c in row.cells for p in c.paragraphs
                 for r in p.runs if r.text.strip()]
    assert cell_runs and all(s == Pt(8) for s in cell_runs)
    corrections = next(p for p in d.paragraphs
                       if p.text == "Corrections along the way")
    assert corrections.paragraph_format.page_break_before is True
    # Other headings do NOT break.
    first = next(p for p in d.paragraphs if p.text.startswith("Provenance"))
    assert not first.paragraph_format.page_break_before
