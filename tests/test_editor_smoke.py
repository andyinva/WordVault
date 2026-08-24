"""
Offscreen smoke test for the stage 2 editor.

Runs the real MainWindow against a temporary database using Qt's
"offscreen" platform plugin, so it works on CI machines and sandboxes
with no display.  Skipped automatically when PyQt6 is not installed
(the storage layer must never require it).
"""

import os

import pytest

# Must be set before Qt is imported anywhere in the process.
os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

PyQt6 = pytest.importorskip("PyQt6.QtWidgets", exc_type=ImportError)

from PyQt6.QtWidgets import QApplication  # noqa: E402

from wordvault import DocumentStore  # noqa: E402
from wordvault.editor import MainWindow  # noqa: E402


@pytest.fixture(scope="module")
def qapp():
    """One QApplication for the whole test module (Qt allows only one)."""
    app = QApplication.instance() or QApplication([])
    yield app


def test_editor_saves_revisions(qapp, tmp_path):
    db = tmp_path / "smoke.db"
    window = MainWindow(db)

    # Create a document through the store (bypassing the title dialog,
    # which would block a headless test), then open it as the UI would.
    doc = window._store.create_document("Smoke Test Essay")
    window._reload_document_list()
    window._open_document(doc.id)

    # Simulate the author typing, then the pause-save firing.
    window._editor.setPlainText("In the beginning was the Word.")
    window._autosave()

    # And an edit followed by another save.
    window._editor.setPlainText(
        "In the beginning was the Word, and the Word was with God."
    )
    window._autosave()

    # Saving identical text must not create a third revision.
    window._autosave()

    revs = window._store.list_revisions(doc.id)
    assert len(revs) == 2
    assert window._store.current_text(doc.id).endswith("with God.")

    # closeEvent path: window closes cleanly and the data survives reopen.
    window.close()
    with DocumentStore(db) as store:
        assert store.current_text(doc.id).endswith("with God.")


def test_loading_a_document_is_not_an_edit(qapp, tmp_path):
    window = MainWindow(tmp_path / "quiet.db")
    doc = window._store.create_document("Quiet Load")
    window._store.save_revision(doc.id, "existing text\n")

    window._open_document(doc.id)
    # Loading used set_text_quietly, so no pause timer should be pending;
    # an immediate autosave must find nothing new to record.
    window._autosave()
    assert len(window._store.list_revisions(doc.id)) == 1
    window.close()


# -- stage 3: time travel ----------------------------------------------------

@pytest.fixture()
def window_with_history(qapp, tmp_path):
    """A window on a document with three known states, open and live."""
    window = MainWindow(tmp_path / "travel.db")
    doc = window._store.create_document("Travel")
    for text in ["state 0\n", "state 1\n", "state 2\n"]:
        window._store.save_revision(doc.id, text)
    window._reload_document_list()
    window._open_document(doc.id)
    yield window, doc
    window.close()


def test_slider_walks_history_read_only(window_with_history):
    window, doc = window_with_history
    assert window._is_live
    assert window._editor.toPlainText() == "state 2\n"

    # Drag the slider back to the oldest revision (as Alt+Left would step).
    window._timeline._slider.setValue(0)
    assert window._editor.toPlainText() == "state 0\n"
    assert window._editor.isReadOnly()      # history is view-only
    assert not window._is_live

    # And forward again to the newest: editable once more.
    window._timeline.go_newest()
    assert window._editor.toPlainText() == "state 2\n"
    assert not window._editor.isReadOnly()
    assert window._is_live


def test_history_mode_is_marked_and_copyable(window_with_history):
    """The time-travel trap (Aug 2026): clicking into an old version
    showed no cursor and no way back.  Now history mode must (a) keep
    the text selectable WITH a keyboard cursor so copying works, (b)
    wear the amber 'history' border, and (c) light up the timeline's
    Newest/Restore buttons — the road back to editing."""
    from PyQt6.QtCore import Qt

    window, doc = window_with_history
    window._timeline._slider.setValue(0)          # into the past

    flags = window._editor.textInteractionFlags()
    assert flags & Qt.TextInteractionFlag.TextSelectableByMouse
    assert flags & Qt.TextInteractionFlag.TextSelectableByKeyboard
    assert window._editor.property("mode") == "history"
    assert window._timeline._newest_btn.styleSheet() != ""
    assert window._timeline._restore_btn.styleSheet() != ""

    window._timeline.go_newest()                  # and back to now
    assert window._editor.property("mode") == "live"
    assert not window._editor.isReadOnly()
    assert window._timeline._newest_btn.styleSheet() == ""
    flags = window._editor.textInteractionFlags()
    assert flags & Qt.TextInteractionFlag.TextEditable


def test_history_stepping_holds_the_readers_place(qapp, tmp_path):
    """The view rules for time travel (Aug 2026, revised twice at the
    author's request): stepping into history HOLDS the reader's place
    (content-anchored — a raw scroll value lies across revisions);
    sitting at the END keeps following the end (the growing edge); and
    Newest restores the departure photograph exactly."""
    import time

    window = MainWindow(tmp_path / "scroll.db")
    doc = window._store.create_document("Long")
    lines = "\n".join(f"line {i}" for i in range(400))
    window._store.save_revision(doc.id, lines + "\nrev one\n")
    window._store.save_revision(doc.id, lines + "\nrev two\n")
    window._store.save_revision(doc.id, lines + "\nrev three\n")
    window._reload_document_list()
    window._open_document(doc.id)
    window.resize(600, 400)

    def pump(ms=400):
        # Restores land on zero-delay timers plus brief retries while
        # the fresh text lays out — run the loop as a real app would.
        end = time.time() + ms / 1000.0
        while time.time() < end:
            qapp.processEvents()
            time.sleep(0.01)

    bar = window._editor.verticalScrollBar()
    if bar.maximum() == 0:
        window.close()
        pytest.skip("offscreen viewport shows the whole document")
    middle = bar.maximum() // 2
    bar.setValue(middle)
    departed = bar.value()

    window._timeline._slider.setValue(1)          # live -> history: HELD
    pump()
    assert abs(bar.value() - departed) <= 3

    window._timeline._slider.setValue(0)          # step again: still held
    pump()
    assert abs(bar.value() - departed) <= 3

    window._timeline._slider.setValue(2)          # Newest: the photograph
    pump()
    assert abs(bar.value() - departed) <= 2

    bar.setValue(bar.maximum())                   # now sit at the end...
    window._timeline._slider.setValue(1)          # ...steps follow the end
    pump()
    assert bar.value() == bar.maximum()
    window.close()


def test_timeline_arrow_buttons_step(window_with_history):
    window, doc = window_with_history
    assert window._is_live
    window._timeline._back_btn.click()
    assert not window._is_live                    # stepped into the past
    assert window._editor.toPlainText() == "state 1\n"
    window._timeline._fwd_btn.click()
    assert window._is_live                        # and back to newest
    assert window._editor.toPlainText() == "state 2\n"


def test_note_anchoring_and_jump_back(qapp, tmp_path):
    """Notes anchored to the text (Aug 2026): the first keystroke of a
    note stamps it with the editor's cursor line and a snippet, and
    the stamp jumps the editor back to that line."""
    from PyQt6.QtCore import QEvent, Qt
    from PyQt6.QtGui import QKeyEvent, QTextCursor

    window = MainWindow(tmp_path / "anchor.db")
    doc = window._store.create_document("Anchored")
    window._store.save_revision(
        doc.id, "first line here\nsecond line words\nthird line text\n")
    window._reload_document_list()
    window._open_document(doc.id)

    # Park the editor's cursor on line 2, then type into the notes.
    block = window._editor.document().findBlockByNumber(1)
    window._editor.setTextCursor(QTextCursor(block))
    press = QKeyEvent(QEvent.Type.KeyPress, Qt.Key.Key_A,
                      Qt.KeyboardModifier.NoModifier, "a")
    QApplication.sendEvent(window._notes, press)

    note = window._notes.toPlainText()
    assert note.startswith("▸ line 2 (second line words): ")
    assert note.endswith("a")                  # the typed letter followed

    # A second keystroke on the SAME line must not stamp again.
    QApplication.sendEvent(window._notes, press)
    assert window._notes.toPlainText().count("▸") == 1

    # The jump: a stamped line moves the editor cursor to its line.
    window._editor.setTextCursor(
        QTextCursor(window._editor.document().firstBlock()))
    assert window._jump_to_note_anchor("▸ line 3 (third): my note")
    assert window._editor.textCursor().blockNumber() == 2
    assert not window._jump_to_note_anchor("an ordinary note line")

    # Sentence anchoring (Aug 2026): a cursor in a LATER sentence of a
    # paragraph stamps that sentence's words, not the paragraph's.
    window._editor.set_text_quietly(
        "One long paragraph. The second thought starts here and "
        "continues on. A third follows.\n")
    doc_block = window._editor.document().firstBlock()
    mid = QTextCursor(doc_block)
    mid.setPosition(doc_block.position() + 40)     # inside sentence two
    window._editor.setTextCursor(mid)
    prefix = window._note_anchor_prefix()
    # The snippet caps at 24 characters, mid-word if need be:
    # "The second thought start" + ellipsis.
    assert prefix.startswith("▸ line 1 (The second thought start")
    assert "…" in prefix
    assert "One long paragraph" not in prefix     # not the paragraph!

    # Drift-proof jump: the stamp's words are FOUND even when the
    # stamped line number no longer holds them.
    window._editor.set_text_quietly(
        "new opening line\nmore new text\n"
        "One long paragraph. The second thought starts here.\n")
    window._editor.setTextCursor(
        QTextCursor(window._editor.document().firstBlock()))
    assert window._jump_to_note_anchor(
        "▸ line 1 (The second thought starts…): my note")
    cursor_now = window._editor.textCursor()
    assert cursor_now.blockNumber() == 2           # found on its NEW line
    rest = window._editor.toPlainText()[cursor_now.position():]
    assert rest.startswith("The second thought starts")

    # Single click: ON the stamp = link; in the note's words = just
    # editing.  The stamp here is 22 characters ("▸ line 3 (third):").
    stamp_line = "▸ line 3 (third): my note"
    assert window._click_is_on_stamp(stamp_line, 0)
    assert window._click_is_on_stamp(stamp_line, 17)     # the colon
    assert not window._click_is_on_stamp(stamp_line, 20)  # "my note"
    assert not window._click_is_on_stamp("plain note", 0)
    window.close()


def test_window_title_carries_version_and_tagline(window_with_history):
    from wordvault import RELEASE_DATE, TAGLINE, __version__

    window, _doc = window_with_history
    title = window.windowTitle()
    assert __version__ in title and RELEASE_DATE in title
    assert TAGLINE in title


def test_title_header_shows_draft_number_and_date(window_with_history):
    window, doc = window_with_history
    text = window._title_label.text()
    assert "draft 3 of 3" in text              # live: the newest state
    window._timeline._slider.setValue(0)       # oldest draft on screen
    assert "draft 1 of 3" in window._title_label.text()
    window._timeline.go_newest()
    assert "draft 3 of 3" in window._title_label.text()


def test_open_external_files_go_straight_into_the_vault(qapp, tmp_path):
    """File > Open (docx/md/txt): convert on open, vault immediately,
    then edit — the document is protected from its first second.  The
    docx path uses the full importer; titles come from the first
    heading (numbered on collision); docx dates come from the file's
    internal record."""
    from datetime import datetime, timezone

    import docx as docx_lib

    window = MainWindow(tmp_path / "open.db")

    # --- markdown: heading becomes the title ---
    md = tmp_path / "05 My Essay.md"
    md.write_text("# My Essay\n\nBody words.\n", encoding="utf-8")
    doc = window._import_external_file(md, "md")
    assert doc.title == "My Essay"
    assert window._store.current_text(doc.id) == "# My Essay\n\nBody words.\n"
    assert doc.original_path == str(md)

    # --- collision: same heading again -> numbered title ---
    doc2 = window._import_external_file(md, "md")
    assert doc2.title == "My Essay (2)"

    # --- txt: plain text, normalized (bytes: write_text on Windows
    # would turn the literal \r\n into \r\r\n and corrupt the test) ---
    txt = tmp_path / "note.txt"
    txt.write_bytes(b"just words\r\nsecond line\r\n")
    doc3 = window._import_external_file(txt, "txt")
    assert doc3.title == "note"
    assert window._store.current_text(doc3.id) == "just words\nsecond line\n"

    # --- docx: full conversion + Word-internal dates ---
    d = docx_lib.Document()
    d.add_paragraph("Opened Chapter", style="Heading 1")
    p = d.add_paragraph()
    p.add_run("with ").bold = False
    p.add_run("weight").bold = True
    d.core_properties.created = datetime(2020, 2, 2, tzinfo=timezone.utc)
    dx = tmp_path / "opened.docx"
    d.save(str(dx))
    doc4 = window._import_external_file(dx, "docx")
    assert doc4.title == "Opened Chapter"
    text = window._store.current_text(doc4.id)
    assert "# Opened Chapter" in text and "**weight**" in text
    assert doc4.created_utc.startswith("2020-02-02")

    # Opening for editing works like any vault document.
    window._open_document(doc4.id)
    assert window._current_doc.id == doc4.id
    assert not window._editor.isReadOnly()
    window.close()


def test_trash_document_from_menu(qapp, tmp_path, monkeypatch):
    """Document > Move to Wastebasket: confirm, banish, editor closes;
    the document leaves the library list but survives underneath."""
    from PyQt6.QtWidgets import QMessageBox

    monkeypatch.setattr(
        QMessageBox, "question",
        staticmethod(lambda *a, **k: QMessageBox.StandardButton.Yes))

    window = MainWindow(tmp_path / "trash.db")
    doc = window._store.create_document("Oops")
    window._store.save_revision(doc.id, "accidental words\n")
    window._reload_document_list()
    window._open_document(doc.id)

    window._on_trash_document()
    assert window._current_doc is None            # editor closed
    assert window._store.list_documents() == []   # gone from the living
    trashed = window._store.list_trashed()
    assert [d.title for d in trashed] == ["Oops"]

    window._store.restore_document(trashed[0].id)
    assert [d.title for d in window._store.list_documents()] == ["Oops"]
    window.close()


def test_notes_spelling_follows_toggle_and_font_applies(qapp, tmp_path):
    """The notes pane checks spelling with the same View toggle as the
    editor, and wears its own Settings font."""
    window = MainWindow(tmp_path / "notesfont.db")
    doc = window._store.create_document("N")
    window._store.save_revision(doc.id, "text\n")
    window._reload_document_list()
    window._open_document(doc.id)

    # The toggle drives BOTH highlighters (when pyspellchecker exists,
    # the action stays checked; without it, both stay off — either
    # way they agree).
    window._spelling_action.setChecked(True)
    assert (window._notes_highlighter.spelling_enabled
            == window._editor.markdown_highlighter.spelling_enabled)
    window._spelling_action.setChecked(False)
    assert window._notes_highlighter.spelling_enabled is False

    # The notes font knobs apply on demand.
    window._settings.setValue("notes_font_pt", 14)
    window._apply_notes_font()
    assert window._notes.font().pointSize() == 14
    window._settings.remove("notes_font_pt")     # tidy the registry
    window.close()


def test_editing_clock_counts_writing_not_open_time(qapp, tmp_path):
    """The editing clock (Aug 2026): keystroke-to-keystroke gaps up to
    a minute count as writing; longer gaps count as absence; flushing
    banks the seconds into the vault under the right document."""
    from wordvault.editor.main_window import _format_edit_time

    window = MainWindow(tmp_path / "clock.db")
    doc = window._store.create_document("Timed")
    window._store.save_revision(doc.id, "start\n")
    window._reload_document_list()
    window._open_document(doc.id)

    # First keystroke: starts the session, credits nothing yet.
    window._on_edit_activity()
    assert window._edit_clock_doc == doc.id
    assert window._edit_pending == 0.0

    # Five seconds of composing between keystrokes: counted.
    window._edit_last_monotonic -= 5.0
    window._on_edit_activity()
    assert 4.5 <= window._edit_pending <= 5.5

    # A twenty-minute absence: NOT counted (open time is not writing).
    window._edit_last_monotonic -= 1200.0
    window._on_edit_activity()
    assert window._edit_pending <= 5.5

    window._flush_edit_clock()
    banked = window._store.editing_seconds(doc.id)
    assert 4 <= banked <= 6
    assert window._edit_pending == 0.0
    # Flushing mid-session does not end the session.
    assert window._edit_clock_doc == doc.id

    # Genuine typing drives the clock through the editor signal too.
    window._editor.insertPlainText("x")
    assert window._edit_clock_doc == doc.id

    # And the human phrasing behaves.
    assert _format_edit_time(0) == "none yet"
    assert _format_edit_time(30) == "under a minute"
    assert _format_edit_time(150) == "2 min"
    assert _format_edit_time(7500) == "2 h 5 min"
    window.close()


def test_delete_selection_leaves_clipboard_alone(qapp, tmp_path):
    from PyQt6.QtGui import QTextCursor
    from PyQt6.QtWidgets import QApplication

    window = MainWindow(tmp_path / "delsel.db")
    doc = window._store.create_document("D")
    window._store.save_revision(doc.id, "keep DELETE keep\n")
    window._reload_document_list()
    window._open_document(doc.id)

    QApplication.clipboard().setText("precious clipboard")
    cursor = window._editor.textCursor()
    cursor.setPosition(5)
    cursor.setPosition(11, QTextCursor.MoveMode.KeepAnchor)
    window._editor.setTextCursor(cursor)
    window._on_delete_selection()
    assert window._editor.toPlainText() == "keep  keep\n"
    assert QApplication.clipboard().text() == "precious clipboard"
    window.close()


def test_dictionary_listing_lets_errors_look_up_their_words():
    """The Aug 2026 request in miniature: typing a misspelling (or
    its start) surfaces 'typed → corrected' from the author's own
    history — errors become index entries pointing at the truth."""
    from wordvault.editor.main_window import _dictionary_listing

    personal = ["habakkuk", "jeremiah"]
    pairs = [("jeprodising", "jeopardizing", 3), ("teh", "the", 1)]
    completions = ["jeopardy", "jeopardize"]

    rows = _dictionary_listing("jep", personal, pairs, completions)
    assert "jeprodising → jeopardizing  (3×)" in rows
    # The CORRECT side works as a key too.
    rows = _dictionary_listing("jeo", personal, pairs, completions)
    assert any(r.startswith("jeprodising →") for r in rows)
    # Personal words lead, marked; standard completions follow.
    rows = _dictionary_listing("je", personal, pairs, completions)
    assert rows[0] == "★ jeremiah"
    assert "jeopardy" in rows
    # Single-occurrence pairs show without a count.
    rows = _dictionary_listing("teh", personal, pairs, [])
    assert rows == ["teh → the"]


def test_speakable_strips_typography():
    """Read Aloud must SAY words, not markup: 'kingdom', never
    'asterisk asterisk kingdom'."""
    from wordvault.editor.main_window import _speakable

    markdown = ("# Chapter One\n\n"
                "The **kingdom** appears *often* here.\n\n"
                "> A quoted verse.\n\n"
                "- first point\n"
                "2. second point\n")
    spoken = _speakable(markdown)
    assert "Chapter One" in spoken and "#" not in spoken
    assert "The kingdom appears often here." in spoken
    assert "A quoted verse." in spoken and ">" not in spoken
    assert "first point" in spoken and "- " not in spoken
    assert "second point" in spoken and "2." not in spoken
    assert "*" not in spoken


def test_sentence_start_finds_the_nearest_sentence():
    """Read Aloud backs up to the SENTENCE, not the paragraph."""
    from wordvault.editor.main_window import _sentence_start

    text = 'First thought here. Second one follows! "Third?" And last.'
    assert _sentence_start(text, 5) == 0            # in the first sentence
    assert _sentence_start(text, 25) == 20          # in "Second one..."
    assert text[_sentence_start(text, 25):].startswith("Second")
    assert text[_sentence_start(text, 45):].startswith('"Third?"')
    assert text[_sentence_start(text, 58):].startswith("And last.")
    # Cursor right at the paragraph's start: sentence one, offset 0.
    assert _sentence_start(text, 0) == 0


def test_speakable_map_points_back_to_the_document():
    """The karaoke map: every spoken character knows its index in the
    original marked-up text, so the engine's word offsets can light
    the right span in the editor."""
    from wordvault.editor.main_window import _speakable_mapped

    original = "# Head\n\nThe **kingdom** here.\n"
    spoken, positions = _speakable_mapped(original, base=100)
    assert len(spoken) == len(positions)
    assert spoken.startswith("Head\n")
    # "Head" maps to its true place after the stripped "# ".
    h = spoken.index("Head")
    assert positions[h] == 100 + original.index("Head")
    # "kingdom" maps back INSIDE the ** marks.
    k = spoken.index("kingdom")
    assert positions[k] == 100 + original.index("kingdom")
    # The characters after the marks keep alignment too.
    w = spoken.index("here")
    assert positions[w] == 100 + original.index("here")


def test_dark_mode_applies_and_reverts(qapp, tmp_path):
    """The Settings checkbox: dark dresses the palette and our own
    surfaces; unchecking truly restores the platform's look."""
    from PyQt6.QtGui import QPalette

    window = MainWindow(tmp_path / "theme.db")
    light_window_color = qapp.palette().color(QPalette.ColorRole.Window)

    window._apply_theme(True)
    dark_color = qapp.palette().color(QPalette.ColorRole.Window)
    assert dark_color.lightness() < 100          # genuinely dark
    assert "#202226" in window._title_label.styleSheet()
    assert "#26251f" in window._notes.styleSheet()
    assert "#2e2a20" in window._editor.styleSheet()   # dark history amber
    # The framed side panels follow the theme (the Outline stayed
    # glaring white in the dark once — never again).
    assert len(window._panel_frames) == 3
    for panel in window._panel_frames:
        assert "#232428" in panel.styleSheet()

    window._apply_theme(False)
    back = qapp.palette().color(QPalette.ColorRole.Window)
    # Light restore must return the PLATFORM'S OWN palette exactly —
    # not a generic look-alike (the flat-gray light-mode bug).
    assert back == light_window_color
    assert "#f4f6f8" in window._title_label.styleSheet()
    for panel in window._panel_frames:          # frames come home too
        assert "#b9c4d0" in panel.styleSheet()
        assert "#232428" not in panel.styleSheet()
    window.close()


def test_read_button_exists(window_with_history):
    window, _doc = window_with_history
    assert window._read_btn.text() == "🔊 Read"


def test_autosave_refuses_in_history_mode(window_with_history):
    # The guard rail: viewing old text must never be saved as new typing.
    window, doc = window_with_history
    window._timeline._slider.setValue(0)    # now viewing "state 0"
    window._autosave()                      # e.g. a stray Ctrl+S
    assert len(window._store.list_revisions(doc.id)) == 3  # unchanged


def test_leaving_live_mode_saves_pending_words(window_with_history):
    # Typing, then dragging the slider back: the unsaved words must be
    # captured as a revision BEFORE the view switches to history.
    window, doc = window_with_history
    window._editor.setPlainText("state 3, not yet auto-saved\n")
    window._timeline._slider.setValue(0)
    texts = [window._store.get_text(r.id)
             for r in window._store.list_revisions(doc.id)]
    assert "state 3, not yet auto-saved\n" in texts


def test_restore_appends_new_revision(window_with_history):
    window, doc = window_with_history
    window._timeline._slider.setValue(1)    # viewing "state 1"
    window._on_restore()

    history = window._store.list_revisions(doc.id)
    assert len(history) == 4                          # appended, not rewritten
    assert history[-1].origin == "restore"
    assert window._store.current_text(doc.id) == "state 1\n"
    assert window._is_live                            # back to editing
    assert not window._editor.isReadOnly()


def test_restore_does_nothing_when_live(window_with_history):
    window, doc = window_with_history
    window._on_restore()                    # Ctrl+R while at the newest
    assert len(window._store.list_revisions(doc.id)) == 3


# -- stage 7: focus mode, age colors, tags ----------------------------------

DOC_TEXT = "# One\nalpha\nbeta\n# Two\ngamma\n"


@pytest.fixture()
def window_with_sections(qapp, tmp_path):
    window = MainWindow(tmp_path / "st7.db")
    doc = window._store.create_document("Sections")
    window._store.save_revision(doc.id, DOC_TEXT)
    window._reload_document_list()
    window._open_document(doc.id)
    yield window, doc
    window.close()


def test_focus_mode_hides_other_sections(window_with_sections):
    window, doc = window_with_sections
    # Cursor into "alpha" (line 1), then hoist.
    editor = window._editor
    cursor = editor.textCursor()
    cursor.setPosition(editor.document().findBlockByNumber(1).position())
    editor.setTextCursor(cursor)
    window._on_focus_section()

    blocks = editor.document()
    visible = [blocks.findBlockByNumber(i).isVisible() for i in range(5)]
    assert visible == [True, True, True, False, False]   # only section One
    assert editor.is_focused()

    window._on_unfocus()
    assert all(blocks.findBlockByNumber(i).isVisible() for i in range(5))


def test_age_colors_toggle_no_crash_and_clear(window_with_sections):
    window, doc = window_with_sections
    # A second revision so there ARE two ages to tint.
    window._editor.setPlainText(DOC_TEXT + "delta added later\n")
    window._autosave()
    window._age_action.setChecked(True)
    assert len(window._editor.extraSelections()) >= 1   # old lines tinted
    window._age_action.setChecked(False)
    assert window._editor.extraSelections() == []


def test_outline_pane_follows_document(window_with_sections):
    window, doc = window_with_sections
    assert window._outline.topLevelItemCount() == 2     # One, Two


def test_close_document_resets_editor(window_with_sections):
    window, doc = window_with_sections
    assert window._current_doc is not None
    window._on_close_document()
    assert window._current_doc is None
    assert not window._editor.isEnabled()
    # Closing again is a harmless no-op.
    window._on_close_document()


def test_line_number_toggle(window_with_sections):
    # NOTE: MainWindow loads the REAL user's persisted QSettings, so the
    # gutter may start on or off — assert the toggle from either state,
    # and restore the user's preference afterwards.
    window, doc = window_with_sections
    initial = window._line_numbers_action.isChecked()
    try:
        window._line_numbers_action.setChecked(True)
        assert window._editor.line_numbers_visible()
        assert window._editor.line_number_width() > 0
        window._line_numbers_action.setChecked(False)
        assert not window._editor.line_numbers_visible()
        assert window._editor.line_number_width() == 0
    finally:
        window._line_numbers_action.setChecked(initial)


def test_library_info_panel_populates(window_with_sections):
    window, doc = window_with_sections
    panel = window._library_panel
    assert panel._documents.text() == "1"
    assert panel._name.text().endswith(".db")
    assert panel._location.text()          # the path line edit is filled
    # A new revision bumps the counts on save.
    before = int(panel._revisions.text().replace(",", ""))
    window._editor.setPlainText(DOC_TEXT + "more\n")
    window._autosave()
    assert int(panel._revisions.text().replace(",", "")) == before + 1


def test_recent_menu_lists_opened_documents(window_with_sections):
    window, doc = window_with_sections
    saved = window._settings.value("recent_docs", [])  # the user's real list
    try:
        window._settings.setValue("recent_docs", [])   # isolate the test
        window._open_document(doc.id)
        window._rebuild_recent_menu()
        titles = [a.text() for a in window._recent_menu.actions()]
        assert "Sections" in titles
    finally:
        window._settings.setValue("recent_docs", saved)


def test_title_header_follows_document(window_with_sections):
    window, doc = window_with_sections
    # The header opens with the title and now carries the draft
    # number and date after it (see the draft-header test).
    assert window._title_label.text().startswith("Sections")
    assert "draft" in window._title_label.text()
    window._on_close_document()
    assert window._title_label.text() == "No document open"


def test_notes_pane_saves_and_travels_per_document(window_with_sections):
    window, doc = window_with_sections
    window._notes.setPlainText("tighten the second section")
    window._save_current_note()

    other = window._store.create_document("Other Doc")
    window._store.save_revision(other.id, "other text\n")
    window._open_document(other.id)                # switch saves + loads
    assert window._notes.toPlainText() == ""       # fresh doc: empty note

    window._notes.setPlainText("note on the other doc")
    window._open_document(doc.id)                  # switching back saves it
    assert window._notes.toPlainText() == "tighten the second section"
    assert window._store.get_note(other.id) == "note on the other doc"


def test_notes_editing_never_creates_revisions(window_with_sections):
    window, doc = window_with_sections
    before = len(window._store.list_revisions(doc.id))
    window._notes.setPlainText("just thinking out loud")
    window._save_current_note()
    assert len(window._store.list_revisions(doc.id)) == before


def test_dock_toggle_actions_exist(window_with_sections):
    window, doc = window_with_sections
    # The author can hide the Library list and the Library Info panels.
    assert window._library_list_dock.toggleViewAction() is not None
    window._library_list_dock.close()
    assert not window._library_list_dock.isVisible()
    window._library_list_dock.toggleViewAction().trigger()


def test_window_state_persists_across_sessions(qapp, tmp_path):
    # Close with the Library list hidden and a document open; a new
    # session must come back the same way.  The user's REAL layout keys
    # are snapshotted and restored — tests must not rearrange their app.
    from PyQt6.QtCore import QSettings
    settings = QSettings("WordVault", "WordVault")
    saved = {k: settings.value(k)
             for k in ("win_geometry", "win_state", "split_state")}
    db = tmp_path / "persist.db"
    try:
        window1 = MainWindow(db)
        doc = window1._store.create_document("Resume Me")
        window1._store.save_revision(doc.id, "where I left off\n")
        window1._reload_document_list()
        window1._open_document(doc.id)
        window1._library_list_dock.hide()
        window1.close()                          # closeEvent saves state

        window2 = MainWindow(db)
        assert window2._library_list_dock.isHidden()
        assert window2._current_doc is not None
        assert window2._current_doc.title == "Resume Me"
        assert window2._title_label.text().startswith("Resume Me")
        window2._library_list_dock.show()        # tidy up for later tests
        window2.close()
    finally:
        for key, value in saved.items():
            if value is None:
                settings.remove(key)
            else:
                settings.setValue(key, value)
        settings.remove(f"last_doc:{db}")


def test_tag_filter_narrows_library(window_with_sections):
    window, doc = window_with_sections
    other = window._store.create_document("Untagged")
    window._store.add_tag(doc.id, "genesis")
    window._reload_tag_filter()
    window._tag_filter.setCurrentText("genesis")

    titles = [window._doc_list.item(i).text()
              for i in range(window._doc_list.count())]
    assert titles == ["Sections"]
    window._tag_filter.setCurrentText("All documents")
    assert window._doc_list.count() == 2
