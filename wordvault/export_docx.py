"""
export_docx.py — Markdown conventions -> a real Word document.

The exact REVERSE of the importer (wordvault/ingest/extract.py):
where ingest turns Word styles into our Markdown, this turns our
Markdown back into Word styles, so a document can leave the vault as
cleanly as it arrived:

    # .. ######    ->  Heading 1..6 styles
    > text        ->  the Quote style (italic fallback if absent)
    - text        ->  List Bullet style
    1. text       ->  List Number style
    **b** *i*     ->  bold / italic / bold-italic runs
    blank line    ->  paragraph break (consecutive plain lines join
                      into one paragraph, per the house conventions)

Standard python-docx (already required by the importer); no other
dependencies.  The Word file also carries title and author in its
core properties — the same fields the importer reads back as a
document's true dates and identity.
"""

from __future__ import annotations

import re
from pathlib import Path

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_RE_BULLET = re.compile(r"^[-*]\s+(.*)$")
_RE_NUMBERED = re.compile(r"^\d{1,3}\.\s+(.*)$")
_RE_QUOTE = re.compile(r"^>\s?(.*)$")
_RE_INLINE = re.compile(
    r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|(?<!\*)\*([^*]+?)\*(?!\*)"
)


def _add_runs(paragraph, text: str) -> None:
    """Write one line's text as runs, honoring **bold** / *italic* /
    ***both*** (unmatched markers stay literal, as everywhere)."""
    pos = 0
    for match in _RE_INLINE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(1))
            run.bold = run.italic = True
        elif match.group(2) is not None:
            paragraph.add_run(match.group(2)).bold = True
        else:
            paragraph.add_run(match.group(3)).italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _styled_paragraph(document, style_name: str):
    """A paragraph in the named style, or plain when the template
    lacks it (older/stripped templates) — never an error."""
    try:
        return document.add_paragraph(style=style_name)
    except KeyError:
        return document.add_paragraph()


def markdown_to_docx(markdown_text: str, path: str | Path, *,
                     title: str = "", author: str = "",
                     compact: bool = False,
                     page_break_before: tuple = ()) -> None:
    """Write `markdown_text` to `path` as a .docx (see module doc).

    compact: the provenance-report dress, learned from Andrew's own
    hand-tuned file — 0.4-inch margins, everything at 8 pt, single
    line spacing, whisper-thin paragraph gaps — so a document's whole
    story fits on one page.  page_break_before: heading texts that
    must start a fresh page (the spelling corrections go to page 2).
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt

    document = DocxDocument()
    if title:
        document.core_properties.title = title
    if author:
        document.core_properties.author = author
    if compact:
        section = document.sections[0]
        for side in ("top_margin", "bottom_margin",
                     "left_margin", "right_margin"):
            setattr(section, side, Inches(0.4))

    def compacted(paragraph, *, before=0.0, after=2.0, size=8.0):
        """Apply the 8-pt single-spaced dress to one paragraph."""
        if not compact:
            return
        pf = paragraph.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(size)

    paragraph_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush() -> None:
        if paragraph_lines:
            paragraph = document.add_paragraph()
            _add_runs(paragraph, " ".join(paragraph_lines))
            compacted(paragraph)
            paragraph_lines.clear()

    def flush_table() -> None:
        """Consecutive |pipe| lines become a REAL Word table (first
        row as its header) — the provenance report's session table
        arrives in Word as a table, not a wall of pipes."""
        if not table_rows:
            return
        columns = max(len(r) for r in table_rows)
        table = document.add_table(rows=len(table_rows), cols=columns)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass                          # stripped template: bare table
        for r, cells in enumerate(table_rows):
            for c in range(columns):
                text = cells[c] if c < len(cells) else ""
                paragraph = table.cell(r, c).paragraphs[0]
                _add_runs(paragraph, text)
                if r == 0:
                    for run in paragraph.runs or [paragraph.add_run("")]:
                        run.bold = True
                compacted(paragraph, after=0.0)
        table_rows.clear()

    for line in markdown_text.split("\n"):
        if line.strip().startswith("|"):
            flush()
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                table_rows.append(cells)  # skip the |---|---| separator
            continue
        flush_table()
        heading = _RE_HEADING.match(line)
        bullet = _RE_BULLET.match(line)
        numbered = _RE_NUMBERED.match(line)
        quote = _RE_QUOTE.match(line)

        if heading:
            flush()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            paragraph = _styled_paragraph(document, f"Heading {level}")
            _add_runs(paragraph, text)
            compacted(paragraph, before=6.0, after=1.0)
            if text in page_break_before:
                paragraph.paragraph_format.page_break_before = True
        elif quote:
            flush()
            paragraph = _styled_paragraph(document, "Quote")
            _add_runs(paragraph, quote.group(1))
            compacted(paragraph)
        elif bullet:
            flush()
            paragraph = _styled_paragraph(document, "List Bullet")
            _add_runs(paragraph, bullet.group(1))
            compacted(paragraph, after=0.5)
        elif numbered:
            flush()
            paragraph = _styled_paragraph(document, "List Number")
            _add_runs(paragraph, numbered.group(1))
            compacted(paragraph, after=0.5)
        elif not line.strip():
            flush()
        else:
            paragraph_lines.append(line.strip())
    flush()
    flush_table()

    document.save(str(path))
