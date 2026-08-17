"""
extract.py — pull plain text out of .docx files (ingest Phase A).

Formatting is discarded BY DESIGN (DESIGN.md section 2): WordVault stores
plain UTF-8 text; styling belongs to the future Formatter app.  We keep
paragraph breaks (one per line) because they carry structure, not style.

Requires python-docx:  pip install python-docx
The import lives inside extract_text() so that the rest of WordVault
(editor, storage) never needs python-docx installed.
"""

from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Union


def long_path(path: Union[str, Path]) -> str:
    """
    Return a path string safe for very long paths on Windows.

    Windows historically limits paths to 260 characters (MAX_PATH); some
    essay filenames are long enough to exceed it, which makes open() and
    os.stat() fail even though the directory listing shows the file.  The
    '\\\\?\\' extended-length prefix on an absolute path lifts the limit.
    On Linux (and for short Windows paths) this returns the path unchanged.
    """
    p = str(Path(path).absolute())
    if os.name == "nt" and len(p) > 240 and not p.startswith("\\\\?\\"):
        p = "\\\\?\\" + p
    return p


def normalize_text(text: str) -> str:
    """
    Normalize extracted text to WordVault's storage form
    (DESIGN.md section 12): LF line endings, no trailing whitespace on
    lines, exactly one trailing newline on non-empty text.

    Runs of blank lines are capped at TWO.  Word documents often contain
    a dozen empty paragraphs in a row (title pages, manual page spacing);
    kept verbatim they become walls of blank space in the editor.  Capped
    at two, the text reads like a hand-written Markdown file: one blank
    line between paragraphs, at most a double break between sections.

    The same normalization is applied before exact-duplicate hashing, so
    two files that differ only in line endings, trailing spaces, or
    blank-line padding are recognized as the same text.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = []
    blank_run = 0
    for line in (l.rstrip() for l in text.split("\n")):
        if line == "":
            blank_run += 1
            if blank_run > 2:
                continue          # cap the run — skip the excess blanks
        else:
            blank_run = 0
        lines.append(line)

    body = "\n".join(lines).strip("\n")
    return body + "\n" if body else ""


def extract_text(path: Union[str, Path]) -> str:
    """
    Extract normalized plain text from one .docx file — every scrap of
    formatting discarded.

    Paragraph text only.  Tables/headers/footnotes are out of scope for
    essays; they can be added here later without touching anything else.
    """
    from docx import Document as DocxDocument  # deferred import, see module doc

    docx = DocxDocument(long_path(path))  # long_path: survive >260-char paths
    return normalize_text("\n".join(p.text for p in docx.paragraphs))


# -- Markdown extraction ------------------------------------------------------
#
# The design keeps documents as plain text, but plain text can CARRY a
# little structure by convention: Markdown.  These mappings translate the
# most meaningful Word formatting into Markdown so it survives ingest and
# can be mapped back to Word styles by the Formatter:
#
#     Word style "Heading 1..6"      ->  # .. ######
#     Word style "Title"/"Subtitle"  ->  # / ##
#     Word styles "Quote"/"Intense Quote"  ->  > blockquote
#     indented plain paragraph (>= 1/3")   ->  > blockquote
#     List styles OR toolbar lists (numPr) ->  - item / 1. item
#     bold / italic runs             ->  **bold** / *italic* / ***both***
#     underlined runs                ->  *italic* (emphasis preserved)
#     hyperlinks                     ->  the text (its address)
#     tables                         ->  one line per row, cells " — "
#
# Everything else (fonts, sizes, colors, alignment) is aesthetics, not
# structure — deliberately dropped, per DESIGN.md section 2.

def _iter_paragraph_parts(paragraph):
    """Yield ('run', bold, italic, text) and ('link', text, url) in
    document order — INCLUDING text inside hyperlinks.  (python-docx's
    paragraph.runs silently omits hyperlink contents, so linked words
    used to vanish on import — genuine text loss, fixed here.)"""
    from docx.oxml.ns import qn
    from docx.text.run import Run

    for child in paragraph._p:
        if child.tag == qn("w:r"):
            run = Run(child, paragraph)
            if run.text:
                # Underline maps to italic: no Markdown underline
                # exists, and underlining is emphasis in manuscripts.
                yield ("run", bool(run.bold),
                       bool(run.italic) or bool(run.underline), run.text)
        elif child.tag == qn("w:hyperlink"):
            text = "".join(Run(r, paragraph).text or ""
                           for r in child.findall(qn("w:r")))
            if not text:
                continue
            url = ""
            rid = child.get(qn("r:id"))
            if rid:       # absent for internal anchors (e.g. TOC links)
                try:
                    url = paragraph.part.rels[rid].target_ref
                except KeyError:
                    url = ""
            yield ("link", text, url)


def _runs_to_markdown(paragraph) -> str:
    """One paragraph's content -> text with **bold** / *italic*
    markers, hyperlinks kept as 'text (address)'.  Adjacent runs with
    identical formatting are merged first, because Word often splits a
    single visually-uniform phrase into many runs."""
    merged: list[list] = []   # [bold, italic, text] — or a link string
    for part in _iter_paragraph_parts(paragraph):
        if part[0] == "link":
            _kind, text, url = part
            rendered = f"{text} ({url})" if url else text
            merged.append([None, None, rendered])   # links stay plain
            continue
        _kind, bold, italic, text = part
        if merged and merged[-1][0] == bold and merged[-1][1] == italic:
            merged[-1][2] += text
        else:
            merged.append([bold, italic, text])

    parts: list[str] = []
    for bold, italic, text in merged:
        if not text.strip() or not (bold or italic):
            parts.append(text)      # plain, link, or whitespace: as-is
            continue
        # Markers must hug the words, not surrounding spaces, or Markdown
        # renderers refuse them ("** bold **" is not bold).
        lead = text[: len(text) - len(text.lstrip())]
        trail = text[len(text.rstrip()):]
        marker = "***" if bold and italic else "**" if bold else "*"
        parts.append(f"{lead}{marker}{text.strip()}{marker}{trail}")
    return "".join(parts)


_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _numbering_formats(path) -> dict:
    """(numId, ilvl) -> 'bullet' | 'number', read straight from the
    document's word/numbering.xml (dependency-free, so it works with
    any python-docx version).  Empty when the file has no lists."""
    import xml.etree.ElementTree as ET
    import zipfile

    try:
        with zipfile.ZipFile(long_path(path)) as zf:
            root = ET.fromstring(zf.read("word/numbering.xml"))
    except (OSError, KeyError, ET.ParseError):
        return {}

    abstract: dict = {}
    for abs_el in root.findall(f"{_W}abstractNum"):
        abs_id = abs_el.get(f"{_W}abstractNumId")
        for lvl in abs_el.findall(f"{_W}lvl"):
            ilvl = int(lvl.get(f"{_W}ilvl", "0"))
            fmt_el = lvl.find(f"{_W}numFmt")
            fmt = fmt_el.get(f"{_W}val") if fmt_el is not None else "bullet"
            abstract[(abs_id, ilvl)] = \
                "bullet" if fmt in ("bullet", "none") else "number"

    formats: dict = {}
    for num_el in root.findall(f"{_W}num"):
        num_id = num_el.get(f"{_W}numId")
        ref = num_el.find(f"{_W}abstractNumId")
        abs_id = ref.get(f"{_W}val") if ref is not None else None
        for (a_id, ilvl), kind in abstract.items():
            if a_id == abs_id and num_id is not None:
                formats[(int(num_id), ilvl)] = kind
    return formats


def _toolbar_list_kind(paragraph, numbering: dict):
    """'bullet' | 'number' for a TOOLBAR-made list item — a numbering
    reference (w:numPr) on an ordinary paragraph.  Most people's lists
    are made this way (the ribbon buttons), not with list STYLES, and
    they used to import as plain paragraphs.  None = not a list."""
    pPr = paragraph._p.pPr
    if pPr is None or pPr.numPr is None or pPr.numPr.numId is None:
        return None
    num_id = pPr.numPr.numId.val
    if num_id in (None, 0):        # numId 0 explicitly REMOVES numbering
        return None
    ilvl = (pPr.numPr.ilvl.val
            if pPr.numPr.ilvl is not None else 0) or 0
    return (numbering.get((num_id, ilvl))
            or numbering.get((num_id, 0))
            or "bullet")           # a list of unknown shape: bullets


def _effective_spacing_pts(paragraph, attribute: str):
    """
    A paragraph-format length ('space_after', 'space_before', or
    'left_indent') in points, honoring style inheritance: the
    paragraph's own setting wins, then its style, then the style's
    base styles.  None = nothing set anywhere we can see.
    """
    value = getattr(paragraph.paragraph_format, attribute)
    if value is not None:
        return value.pt
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        try:
            value = getattr(style.paragraph_format, attribute)
        except AttributeError:
            value = None
        if value is not None:
            return value.pt
        style = style.base_style
    return None


def extract_markdown(path: Union[str, Path]) -> str:
    """
    Extract text from a .docx, translating structural formatting to
    Markdown (see the mapping table above).  This is the default
    extraction for ingest; extract_text() remains for a pure-plain run.

    Paragraph separation is SIMPLE by design (Aug 2026, replacing an
    earlier spacing-faithful heuristic): every paragraph is separated
    from its neighbor by one blank line — the way the editor and the
    print renderer read paragraphs — with two block exceptions:
    consecutive items of one list, lines of one quotation, and rows of
    one table stay tight, as single blocks.

    The body is walked in DOCUMENT ORDER (paragraphs and tables
    interleaved), so table text lands where the table stood.
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    docx = DocxDocument(long_path(path))
    numbering = _numbering_formats(path)
    lines: list[str] = []
    prev_kind: str = ""    # heading | quote | list | table | plain | ''

    def emit(kind: str, line: str) -> None:
        """Append one rendered line under the separation rule."""
        nonlocal prev_kind
        tight = kind == prev_kind and kind in ("list", "quote", "table")
        if lines and not tight:
            lines.append("")
        lines.append(line)
        prev_kind = kind

    for child in docx.element.body.iterchildren():
        if child.tag == qn("w:tbl"):
            # Tables flatten to one line per row (' — ' between cells):
            # the grid is lost, the words are not — they used to be.
            for row in Table(child, docx).rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    emit("table", " — ".join(cells))
            continue
        if child.tag != qn("w:p"):
            continue

        p = Paragraph(child, docx)
        style = (p.style.name if p.style is not None else "") or ""
        style_lower = style.lower()
        toolbar_list = _toolbar_list_kind(p, numbering)

        if style_lower.startswith("heading"):
            digits = "".join(ch for ch in style if ch.isdigit())
            level = min(int(digits), 6) if digits else 1
            emit("heading", "#" * level + " " + p.text.strip())
        elif style_lower == "title":
            emit("heading", "# " + p.text.strip())
        elif style_lower == "subtitle":
            emit("heading", "## " + p.text.strip())
        elif "quote" in style_lower:
            emit("quote", "> " + _runs_to_markdown(p))
        elif style_lower.startswith("list bullet"):
            emit("list", "- " + _runs_to_markdown(p))
        elif style_lower.startswith("list number"):
            emit("list", "1. " + _runs_to_markdown(p))
        elif toolbar_list is not None:
            marker = "- " if toolbar_list == "bullet" else "1. "
            emit("list", marker + _runs_to_markdown(p))
        else:
            text = _runs_to_markdown(p)
            if not text.strip():
                prev_kind = ""     # an empty paragraph ends any block
                continue
            # Indented plain paragraph (>= ~1/3 inch): the common way
            # Scripture is quoted without using the Quote style.
            indent = _effective_spacing_pts(p, "left_indent")
            if indent is not None and indent >= 24:
                emit("quote", "> " + text)
            else:
                emit("plain", text)

    return normalize_text("\n".join(lines))


def file_dates_utc(path: Union[str, Path]) -> tuple[str, str]:
    """
    (created_utc, modified_utc) for a file, as stored ISO-8601 UTC strings.

    On Windows st_ctime is true creation time; on Linux it is metadata-
    change time — so we take min(ctime, mtime) as the best available
    "written when" estimate, and mtime as the modification date.  These
    become the document's created_utc / original_mtime, keeping the
    library ordered by when the material was actually written.
    """
    st = os.stat(long_path(path))  # long_path: survive >260-char paths
    to_iso = lambda ts: datetime.fromtimestamp(ts, timezone.utc).isoformat()
    return to_iso(min(st.st_ctime, st.st_mtime)), to_iso(st.st_mtime)


def docx_internal_dates(path: Union[str, Path]):
    """
    (created_utc, modified_utc) from the docx file's OWN core
    properties (docProps/core.xml) — the dates Word recorded at
    authoring time.  None where a property is absent or unreadable.

    Why these outrank the filesystem: every copy, move, cloud sync, or
    WSL transfer can reset a file's filesystem dates to the day of the
    copy, but the dates inside the document travel with it untouched.
    """
    import xml.etree.ElementTree as ET
    import zipfile

    try:
        with zipfile.ZipFile(long_path(path)) as zf:
            root = ET.fromstring(zf.read("docProps/core.xml"))
    except (OSError, KeyError, ET.ParseError, zipfile.BadZipFile):
        return None, None

    def read(tag: str):
        el = root.find(f"{{http://purl.org/dc/terms/}}{tag}")
        raw = (el.text or "").strip() if el is not None else ""
        if not raw:
            return None
        try:
            moment = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except ValueError:
            return None
        if moment.tzinfo is None:      # bare W3CDTF times are UTC
            moment = moment.replace(tzinfo=timezone.utc)
        return moment.astimezone(timezone.utc).isoformat()

    return read("created"), read("modified")


def document_dates_utc(path: Union[str, Path]) -> tuple[str, str]:
    """The best available (created_utc, modified_utc) for a document
    file: the Word-internal dates when present, the filesystem's as
    the fallback.  This is what ingest and the date-repair pass use."""
    internal_created, internal_modified = docx_internal_dates(path)
    fs_created, fs_modified = file_dates_utc(path)
    return internal_created or fs_created, internal_modified or fs_modified
